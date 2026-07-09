"""DOCX loader (lazy python-docx import)."""

from __future__ import annotations

from pathlib import Path

from ingestion.loaders.base import Document


def load_docx(path: str | Path) -> Document:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - optional extra
        raise RuntimeError(
            'DOCX support needs python-docx. Install with: pip install -e ".[ingest]"'
        ) from exc
    p = Path(path)
    document = docx.Document(str(p))
    text = "\n".join(para.text for para in document.paragraphs)
    return Document(text=text, source=str(p), type="docx")
